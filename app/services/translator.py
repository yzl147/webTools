import os
import uuid
from pathlib import Path
from typing import Optional, Tuple

from app.services.converter import converter_service


class TranslatorService:
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def translate_file(
        self,
        file_path: str,
        file_extension: str,
        mode: str,
        target_language: str,
        lt_host: Optional[str] = None,
        lt_port: Optional[int] = None,
        lt_api_key: Optional[str] = None,
        llm_type: Optional[str] = None,
        api_base: Optional[str] = None,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
    ) -> Tuple[Optional[str], str]:
        """
        Translate file content and return translated file path.
        Returns (file_path, message).
        """
        try:
            content = self._extract_text(file_path, file_extension)
            if not content:
                return None, "Failed to extract text from file"

            if mode == "libretranslate":
                translated = self._translate_libretranslate(
                    content, target_language, lt_host, lt_port, lt_api_key
                )
            elif mode == "llm":
                translated = self._translate_llm(
                    content, target_language, llm_type, api_base, api_key, model
                )
            else:
                return None, f"Unknown translation mode: {mode}"

            if not translated:
                return None, "Translation failed"

            output_name = f"translated_{uuid.uuid4().hex[:8]}.{file_extension}"
            output_path = self.output_dir / output_name

            self._create_file(output_path, translated, file_extension)

            return str(output_path), "Translation completed"

        except Exception as e:
            return None, f"Translation error: {str(e)}"

    def _extract_text(self, file_path: str, file_extension: str) -> Optional[str]:
        """Extract text content from file."""
        try:
            if file_extension == "txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_extension in ["docx"]:
                from docx import Document
                doc = Document(file_path)
                return "\n".join([p.text for p in doc.paragraphs])
            elif file_extension == "pdf":
                import fitz
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
            else:
                temp_pdf = converter_service.convert(file_path, "pdf")
                if temp_pdf:
                    text = self._extract_text(temp_pdf, "pdf")
                    os.unlink(temp_pdf)
                    return text
        except Exception as e:
            print(f"Text extraction error: {e}")
        return None

    def _create_file(self, output_path: Path, content: str, file_extension: str):
        """Create file with translated content."""
        if file_extension == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
        elif file_extension == "docx":
            from docx import Document
            doc = Document()
            for para in content.split("\n"):
                doc.add_paragraph(para)
            doc.save(output_path)
        elif file_extension == "pdf":
            import fitz
            doc = fitz.open()
            for i, line in enumerate(content.split("\n")):
                doc.new_page()
                doc[i].insert_text((72, 72), line, fontsize=11)
            doc.save(output_path)

    def _translate_libretranslate(
        self,
        text: str,
        target_language: str,
        host: Optional[str],
        port: Optional[int],
        api_key: Optional[str],
    ) -> Optional[str]:
        """Translate using LibreTranslate."""
        try:
            import libretranslatepy

            if host and port:
                lt = libretranslatepy.LibreTranslateApi(
                    url=f"http://{host}:{port}",
                    api_key=api_key
                )
            else:
                lt = libretranslatepy.LibreTranslateApi(api_key=api_key)

            result = lt.translate(text, target_lang=target_language)
            return result
        except Exception as e:
            print(f"LibreTranslate error: {e}")
            return None

    def _translate_llm(
        self,
        text: str,
        target_language: str,
        llm_type: Optional[str],
        api_base: Optional[str],
        api_key: Optional[str],
        model: Optional[str],
    ) -> Optional[str]:
        """Translate using LLM (OpenAI or Anthropic)."""
        if not api_key or not model:
            return None

        prompt = f"Translate the following text to {target_language}. Only output the translated text, no explanations:\n\n{text}"

        try:
            if llm_type == "openai":
                return self._translate_openai(prompt, api_base, api_key, model)
            elif llm_type == "anthropic":
                return self._translate_anthropic(prompt, api_key, model)
            else:
                return None
        except Exception as e:
            print(f"LLM translation error: {e}")
            return None

    def _translate_openai(
        self, prompt: str, api_base: Optional[str], api_key: str, model: str
    ) -> Optional[str]:
        """Translate using OpenAI API."""
        from openai import OpenAI

        client = OpenAI(api_key=api_key, base_url=api_base)

        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        return response.choices[0].message.content

    def _translate_anthropic(self, prompt: str, api_key: str, model: str) -> Optional[str]:
        """Translate using Anthropic Claude API."""
        import anthropic

        client = anthropic.Anthropic(api_key=api_key)

        response = client.messages.create(
            model=model,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text


translator_service = TranslatorService()
